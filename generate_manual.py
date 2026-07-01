#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成基因表达调控分析与预测平台软件使用说明书（Word格式）
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name="黑体", size=12):
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(doc, text, level=1):
    heading = doc.add_heading('', level=level)
    run = heading.runs[0] if heading.runs else heading.add_run()
    run.text = text
    set_chinese_font(run, "黑体", 16 if level == 1 else 14 if level == 2 else 12)
    run.bold = True
    return heading


def add_paragraph(doc, text, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 12)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_list_item(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(24 * level)
    return p


def add_numbered_item(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(24 * level)
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.line_spacing = 1.2
    # 设置中文字体备用
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        set_chinese_font(run, "黑体", 11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_chinese_font(run, "宋体", 10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    return table


def add_bold_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    set_chinese_font(run, "宋体", 12)
    p.paragraph_format.line_spacing = 1.5
    return p


def generate_manual():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ===== 封面 =====
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('基因表达调控分析与预测平台')
    run.bold = True
    set_chinese_font(run, "黑体", 28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('软件使用说明书')
    run.bold = True
    set_chinese_font(run, "黑体", 22)

    doc.add_paragraph()
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('V1.0')
    set_chinese_font(run, "宋体", 16)

    for _ in range(8):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('GeneAnalysis Team\n2024年1月')
    set_chinese_font(run, "宋体", 14)

    doc.add_page_break()

    # ===== 目录页 =====
    add_heading(doc, '目录', level=1)
    toc_items = [
        '第一章 前言',
        '  1.1 编写目的',
        '  1.2 项目背景',
        '  1.3 适用对象',
        '第二章 系统架构设计',
        '  2.1 系统功能架构',
        '  2.2 系统数据结构设计',
        '第三章 平台系统操作描述',
        '  3.1 仪表板界面（Web首页）',
        '  3.2 数据配置页面',
        '  3.3 运行分析页面',
        '  3.4 结果展示页面',
        '  3.5 下载结果页面',
        '  3.6 侧边栏导航',
        '  3.7 命令行模式操作',
        '第四章 系统测试',
        '  4.1 功能测试',
        '  4.2 性能测试',
        '第五章 常见问题与解决方案',
        '  5.1 数据上传问题',
        '  5.2 分析参数问题',
        '  5.3 可视化图表问题',
        '  5.4 结果下载问题',
        '  5.5 命令行模式问题',
        '  5.6 Web应用问题',
        '第六章 附录',
        '  6.1 示例数据格式说明',
        '  6.2 参数调整建议参考表',
        '  6.3 技术支持与联系方式',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_chinese_font(run, "宋体", 12)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # ===== 第一章 前言 =====
    add_heading(doc, '第一章 前言', level=1)

    add_heading(doc, '1.1 编写目的', level=2)
    add_paragraph(doc, '本使用说明书旨在为用户提供基因表达调控分析与预测平台的详细操作指南，帮助用户快速掌握平台各项功能的使用方法，规范操作流程，确保分析结果的准确性与可靠性。本说明书适用于生物信息学研究人员、医学研究者、药物开发人员及相关领域的科研工作者。')

    add_heading(doc, '1.2 项目背景', level=2)
    add_paragraph(doc, '随着高通量测序技术的快速发展，基因表达数据呈爆发式增长。如何高效地从海量数据中挖掘有价值的信息，识别差异表达基因，理解基因调控机制，成为生物医学研究的关键问题。传统的分析方法往往需要复杂的编程技能和统计学知识，门槛较高。')
    add_paragraph(doc, '本平台集成了序列分析、差异表达分析、可视化图表生成和机器学习预测四大核心功能，采用Web界面和命令行双入口设计，降低使用门槛，为研究人员提供从数据读取到结果输出的一站式解决方案。平台支持FASTA格式序列文件和CSV格式表达矩阵的导入，提供GC含量计算、核苷酸频率统计、t检验统计分析、BH-FDR多重校正、火山图热图PCA图生成、随机森林基因分类等完整分析流程。')

    add_heading(doc, '1.3 适用对象', level=2)
    add_paragraph(doc, '本平台适用于以下用户群体：')
    add_list_item(doc, '生物信息学研究者：进行基因组特征分析、物种鉴定、序列统计计算')
    add_list_item(doc, '医学研究人员：开展疾病相关基因的差异表达分析、标志物筛选')
    add_list_item(doc, '药物研发人员：分析药物处理后的基因表达变化、识别靶基因')
    add_list_item(doc, '高校师生：学习生物信息学分析方法、完成毕业设计和课题研究')
    add_list_item(doc, '科研院所工作人员：进行高通量数据分析、撰写科研报告')

    doc.add_page_break()

    # ===== 第二章 系统架构设计 =====
    add_heading(doc, '第二章 系统架构设计', level=1)

    add_heading(doc, '2.1 系统功能架构', level=2)
    add_paragraph(doc, '本平台采用模块化架构设计，将复杂分析流程分解为6个独立功能模块，各模块职责清晰，通过标准接口实现数据流转与协同作业。系统架构遵循"应用层-业务逻辑层-功能模块层-基础工具层"四层设计模式，保障系统扩展性与维护便利性。')

    add_bold_text(doc, '数据处理模块')
    add_paragraph(doc, '作为系统入口，集成Biopython库实现FASTA格式序列文件的标准化解析，支持多行序列自动合并、标题行智能识别、序列ID提取等功能，同时提供数据清洗工具，包括重复数据去除、缺失值填充（均值/中位数/零值/向前填充）、数据标准化（Z-score/Min-Max）及异常值过滤，确保输入数据质量符合后续分析要求。该模块采用内存高效的迭代器模式读取，避免大文件一次性加载导致的内存溢出问题，同时支持数据集分类存储与历史记录追溯，保障数据完整性与可复用性。')

    add_bold_text(doc, '序列分析模块')
    add_paragraph(doc, '基于NumPy库，提供GC含量计算、核苷酸频率统计、GC偏斜分析等序列特征计算功能。GC含量计算采用滑动窗口策略，窗口大小可配置（50-500bp），支持捕捉序列局部特征。核苷酸频率统计覆盖A/T/G/C/N五种碱基，自动计算总体频率分布。GC偏斜计算采用(G-C)/(G+C)公式，取值范围[-1,1]，可用于细菌基因组复制起始位点预测。所有计算结果自动生成汇总统计报告，包含平均值、标准差、最大值、最小值等多维度指标。')

    add_bold_text(doc, '差异表达分析模块')
    add_paragraph(doc, '是RNA-seq数据分析的核心组件，集成SciPy统计模块，对每个基因执行独立样本t检验，计算原始P值与log2倍数变化。模块内置Benjamini-Hochberg FDR多重检验校正算法，有效控制假发现率，避免大规模并行检验导致的假阳性膨胀。用户可配置P值阈值（0.01-0.1）与log2FC阈值（0.5-2.0），系统同时满足统计显著性与生物学显著性双重标准，自动标记上调、下调与无变化基因，输出完整差异表达结果表格。')

    add_bold_text(doc, '可视化分析模块')
    add_paragraph(doc, '融合Matplotlib与Seaborn双引擎，生成火山图、表达热图、PCA主成分分析图、GC含量分布图四种核心图表。火山图采用双轴设计，X轴展示log2倍数变化，Y轴展示-log10(校正P值)，通过颜色区分显著上调（红色）、显著下调（蓝色）与无变化基因（灰色）。热图采用Z-score标准化，使不同基因的表达量在同一尺度可比，关注表达模式而非绝对值。PCA图实现样本相似性可视化，相似样本自动聚集，便于检测批次效应与离群样本。所有图表支持300 DPI高清导出，满足论文配图需求。')

    add_bold_text(doc, '机器学习模块')
    add_paragraph(doc, '集成scikit-learn库，采用随机森林（Random Forest）算法对基因进行分类预测。模块支持特征自动选择、数据标准化、训练集/测试集分割、模型训练与评估完整流程。训练过程自动输出准确率、精确率、召回率、F1分数四项评估指标，并提供特征重要性分析，量化各特征对分类结果的贡献度。模型训练完成后支持新样本预测，输出预测标签与概率值，便于科研人员筛选候选基因。')

    add_bold_text(doc, '项目管理与配置模块')
    add_paragraph(doc, '采用dataclass设计模式，集中管理所有分析参数，包括文件路径、统计阈值、机器学习超参数、可视化选项等。支持配置保存与加载，便于实验记录与结果复现。Web界面提供参数实时调整功能，用户可通过滑块、输入框等交互组件灵活配置分析参数。')

    add_heading(doc, '2.2 系统数据结构设计', level=2)
    add_paragraph(doc, '本系统采用面向对象的数据结构设计，基于Python原生数据类型与Pandas库构建核心数据体系，围绕"序列-表达矩阵-差异结果-模型"四层逻辑架构实现全流程数据关联，所有实体通过标准格式建立映射，确保科研数据链路的完整性与可追溯性。')

    add_bold_text(doc, '序列数据结构')
    add_paragraph(doc, '以字典列表形式存储，每个序列字典包含ID、描述、序列字符串、长度四个核心字段，支持FASTAReader类读取后的直接传递。序列ID作为唯一标识符，用于后续分析与结果关联。序列字符串统一转为大写格式，确保碱基计数的准确性。长度字段自动计算，便于统计分析。')

    add_bold_text(doc, '表达矩阵数据结构')
    add_paragraph(doc, '采用Pandas DataFrame构建，行索引为基因名，列名为样本名，单元格存储表达量数值。矩阵支持CSV格式导入，第一列自动识别为基因名索引。数据清洗后自动去除低表达基因，默认阈值平均表达量>10，减少假阳性干扰。样本分组信息通过列名列表形式存储，区分对照组与处理组。')

    add_bold_text(doc, '差异表达结果数据结构')
    add_paragraph(doc, '同样采用DataFrame存储，包含基因名、log2倍数变化、原始P值、校正P值、t统计量、对照组均值、处理组均值、显著性标记、调控方向等字段。显著性标记为布尔值，调控方向为字符串（up/down/no_change）。结果按校正P值升序排序，显著基因排在前列，便于用户快速定位关键基因。')

    add_bold_text(doc, '机器学习模型数据结构')
    add_paragraph(doc, '混合存储模式设计，通过字典封装算法类型、超参数配置、训练指标、特征重要性等元数据，核心模型对象以scikit-learn标准格式持久化。训练集与测试集分割比例、随机种子、决策树数量等参数可配置保存，确保实验可复现。')

    add_bold_text(doc, '可视化图表数据结构')
    add_paragraph(doc, '以matplotlib Figure对象形式存储，支持PNG格式导出。图表元数据包含标题、尺寸、颜色方案、输出路径等信息。Web界面下图表以Base64编码字符串形式传递，便于页面间共享与实时显示。')

    doc.add_page_break()

    # ===== 第三章 平台系统操作描述 =====
    add_heading(doc, '第三章 平台系统操作描述', level=1)

    add_heading(doc, '3.1 仪表板界面（Web首页）', level=2)
    add_paragraph(doc, 'Web应用首页界面作为用户进入平台的第一窗口，展示平台整体功能概览与使用流程引导。界面顶部显示平台名称"基因表达调控分析与预测平台"，主体区域分为欢迎区域、核心功能卡片区域、使用流程表格区域三个部分。')

    add_bold_text(doc, '欢迎区域')
    add_paragraph(doc, '简要介绍平台定位与目标用户群体，说明平台集成的四大核心功能模块，强调双入口设计（Web界面与命令行）与模块化架构特点。')

    add_bold_text(doc, '核心功能卡片区域')
    add_paragraph(doc, '采用四列布局，每列展示一个核心功能模块的简要说明：')
    add_list_item(doc, '第1列：序列分析（GC含量计算、核苷酸频率统计、序列长度分布）')
    add_list_item(doc, '第2列：差异表达分析（t检验统计分析、BH-FDR多重校正、上调/下调基因识别）')
    add_list_item(doc, '第3列：可视化图表（火山图、表达热图、PCA主成分分析、GC含量分布图）')
    add_list_item(doc, '第4列：机器学习（随机森林分类、特征重要性分析、模型性能评估）')

    add_bold_text(doc, '使用流程表格区域')
    add_paragraph(doc, '以表格形式展示四步操作流程：')
    add_table(doc,
        ['步骤', '操作', '说明'],
        [
            ['1', '数据配置', '上传FASTA序列文件和表达矩阵CSV，设置分析参数'],
            ['2', '运行分析', '点击开始按钮，系统自动执行完整分析流程'],
            ['3', '查看结果', '在结果页面浏览各类图表和统计数据'],
            ['4', '下载结果', '导出分析结果表格和高质量图表'],
        ]
    )
    add_paragraph(doc, '界面底部提供快速体验提示，说明系统内置示例数据，用户可直接选择"使用示例数据"快速体验完整分析流程。')

    add_heading(doc, '3.2 数据配置页面', level=2)
    add_paragraph(doc, '数据配置页面是用户使用平台的第一步操作界面，分为数据上传、分析参数设置、样本分组配置三个板块。')

    add_bold_text(doc, '(1) 数据上传板块')
    add_paragraph(doc, '数据上传板块提供两种数据导入方式：')

    add_bold_text(doc, '方式一：使用示例数据')
    add_paragraph(doc, '界面顶部提供"使用内置示例数据（快速体验）"复选框，勾选后自动加载系统内置的模拟数据：')
    add_list_item(doc, 'FASTA序列文件：data/simulated_genes.fasta（约100条模拟基因序列）')
    add_list_item(doc, '表达矩阵文件：data/simulated_expression.csv（约100个基因，6个样本）')
    add_paragraph(doc, '加载成功后界面显示序列数量信息、表达矩阵维度信息、样本名称列表。')

    add_bold_text(doc, '方式二：上传自定义数据')
    add_paragraph(doc, '取消勾选示例数据后，界面显示两个文件上传组件：')
    add_paragraph(doc, '左侧上传组件：FASTA序列文件上传')
    add_list_item(doc, '支持格式：.fasta, .fa, .txt')
    add_list_item(doc, '上传后自动解析序列数量、序列ID、序列长度')
    add_list_item(doc, '显示前3条序列的基本信息')
    add_list_item(doc, '错误提示：文件格式不正确、文件过大、编码问题等')

    add_paragraph(doc, '右侧上传组件：表达矩阵上传')
    add_list_item(doc, '支持格式：.csv')
    add_list_item(doc, '要求：第一列为基因名，其他列为样本表达量')
    add_list_item(doc, '上传后显示基因数量、样本数量、样本名称')
    add_list_item(doc, '自动识别列名，提示数据维度')

    add_bold_text(doc, '(2) 分析参数设置板块')
    add_paragraph(doc, '分析参数设置板块采用两列布局，左侧为差异表达分析参数，右侧为机器学习参数。')

    add_bold_text(doc, '左侧：差异表达分析参数')
    add_list_item(doc, 'P值阈值滑块：范围0.001-0.1，默认值0.05，步长0.001')
    add_list_item(doc, '说明：校正后P值小于此阈值的基因被认为是显著差异基因')
    add_list_item(doc, 'log2倍数变化阈值滑块：范围0.0-3.0，默认值1.0，步长0.1')
    add_list_item(doc, '说明：|log2FC|大于此阈值的基因被认为有生物学意义的变化，1.0对应2倍变化')
    add_list_item(doc, 'GC窗口大小输入框：范围10-500，默认值100，步长10')
    add_list_item(doc, '说明：滑动窗口大小用于计算GC含量分布')

    add_bold_text(doc, '右侧：机器学习参数')
    add_list_item(doc, '测试集比例滑块：范围0.1-0.5，默认值0.2，步长0.05')
    add_list_item(doc, '说明：用于测试模型性能的数据比例')
    add_list_item(doc, '决策树数量滑块：范围10-500，默认值100，步长10')
    add_list_item(doc, '说明：随机森林中决策树的数量，树越多结果越稳定')
    add_list_item(doc, '随机种子输入框：范围0-9999，默认值42，步长1')
    add_list_item(doc, '说明：固定随机种子可以保证结果可复现')

    add_bold_text(doc, '(3) 样本分组配置板块')
    add_paragraph(doc, '样本分组配置板块仅在表达数据上传成功后显示。界面首先展示检测到的样本信息。采用两列布局配置样本分组：')

    add_bold_text(doc, '左侧：对照组样本选择')
    add_list_item(doc, '多选组件，选项为所有样本名')
    add_list_item(doc, '自动猜测：列名包含"control"的样本默认选中')
    add_list_item(doc, '用户可手动调整选择')

    add_bold_text(doc, '右侧：处理组样本选择')
    add_list_item(doc, '多选组件，选项为所有样本名')
    add_list_item(doc, '自动猜测：不在对照组的样本默认选中')
    add_list_item(doc, '用户可手动调整选择')

    add_paragraph(doc, '分组验证：至少每组选择1个样本，否则提示警告；两组不能有重叠样本，否则提示错误；配置正确显示"分组配置完成：对照组X个，处理组X个"。')
    add_paragraph(doc, '界面底部提供"保存配置并继续"按钮，点击后验证配置完整性，通过后提示"配置已保存！请前往【运行分析】页面开始分析"。')

    add_heading(doc, '3.3 运行分析页面', level=2)
    add_paragraph(doc, '运行分析页面是执行完整分析流程的核心界面，分为配置摘要展示区、开始分析按钮区、进度条显示区、实时日志输出区四个部分。')

    add_bold_text(doc, '(1) 配置摘要展示区')
    add_paragraph(doc, '配置摘要以可折叠面板形式展示当前配置信息，默认展开状态。采用三列布局显示：')
    add_list_item(doc, '第一列（数据信息）：序列数量、基因数量、样本数量')
    add_list_item(doc, '第二列（分析参数）：P值阈值、log2FC阈值、GC窗口大小')
    add_list_item(doc, '第三列（样本分组）：对照组数量、处理组数量')

    add_bold_text(doc, '(2) 开始分析按钮区')
    add_paragraph(doc, '界面中部提供分析启动按钮。按钮状态根据分析完成情况动态变化：')
    add_list_item(doc, '未完成状态：显示"开始完整分析流程"按钮（蓝色，主要按钮样式），点击后执行完整分析流程')
    add_list_item(doc, '已完成状态：显示"分析已完成！可以在【结果展示】页面查看结果"提示，提供"重新运行分析"按钮（灰色，次要按钮样式），点击后重置所有状态')

    add_bold_text(doc, '(3) 进度条显示区')
    add_paragraph(doc, '进度条组件显示当前分析进度，值范围0-100（0%到100%）。进度条上方显示当前步骤文本提示：')
    add_list_item(doc, '步骤1/4："序列基础统计分析..."，进度10%')
    add_list_item(doc, '步骤2/4："差异表达分析..."，进度35%')
    add_list_item(doc, '步骤3/4："生成可视化图表..."，进度65%')
    add_list_item(doc, '步骤4/4："机器学习基因分类..."，进度90%')
    add_list_item(doc, '完成："分析完成！"，进度100%')

    add_bold_text(doc, '(4) 实时日志输出区')
    add_paragraph(doc, '日志输出区实时显示分析过程的详细信息，每条日志根据级别使用不同颜色显示：info级别（蓝色）、success级别（绿色）、warning级别（黄色）、error级别（红色）。')

    add_heading(doc, '3.4 结果展示页面', level=2)
    add_paragraph(doc, '结果展示页面是查看分析成果的主要界面，分为统计摘要卡片区、可视化图表标签页区、差异表达结果表格区、机器学习结果展示区四个部分。')

    add_bold_text(doc, '(1) 统计摘要卡片区')
    add_paragraph(doc, '界面顶部采用四列布局，每列显示一个核心指标：')
    add_list_item(doc, '第一列：序列数量指标卡（显示大字号数字，标签"序列数量"）')
    add_list_item(doc, '第二列：平均GC含量指标卡（显示大字号百分比，标签"平均GC含量"）')
    add_list_item(doc, '第三列：显著差异基因数指标卡（显示大字号数字，标签"显著差异基因"）')
    add_list_item(doc, '第四列：模型准确率指标卡（显示大字号百分比，标签"模型准确率"）')

    add_bold_text(doc, '(2) 可视化图表标签页区')
    add_paragraph(doc, '图表区采用标签页（Tabs）设计，用户可切换查看不同图表：')

    add_bold_text(doc, '标签页1：火山图')
    add_list_item(doc, '显示差异表达火山图')
    add_list_item(doc, 'X轴：log2倍数变化（正值=上调，负值=下调）')
    add_list_item(doc, 'Y轴：-log10(P值)（越高越显著）')
    add_list_item(doc, '红色点：显著上调基因，蓝色点：显著下调基因，灰色点：无显著变化基因')
    add_list_item(doc, '支持全屏查看与交互操作')

    add_bold_text(doc, '标签页2：表达热图')
    add_list_item(doc, '显示基因表达热图')
    add_list_item(doc, '每一行是一个基因，每一列是一个样本')
    add_list_item(doc, '颜色深浅表示表达量高低（红色=高表达，蓝色=低表达）')
    add_list_item(doc, '基因和样本都经过聚类，表达模式相似的会聚在一起')
    add_list_item(doc, '基因数量>50时自动隐藏行标签避免拥挤')

    add_bold_text(doc, '标签页3：PCA图')
    add_list_item(doc, '显示样本PCA主成分分析散点图')
    add_list_item(doc, '每个点代表一个样本，点之间的距离表示样本间的相似性')
    add_list_item(doc, '如果对照组和处理组明显分开，说明处理对基因表达有显著影响')
    add_list_item(doc, '轴标签显示主成分解释方差比例')

    add_bold_text(doc, '标签页4：GC分布图')
    add_list_item(doc, '显示序列GC含量分布直方图')
    add_list_item(doc, '展示所有序列的GC含量分布情况')
    add_list_item(doc, 'GC含量影响DNA稳定性和基因表达调控')
    add_list_item(doc, '直方图叠加核密度估计曲线，红色虚线标注均值，绿色虚线标注中位数')

    add_bold_text(doc, '(3) 差异表达结果表格区')
    add_paragraph(doc, '表格区提供差异表达结果的筛选、排序与显示功能。界面顶部三列布局提供筛选控件：')
    add_list_item(doc, '第一列：筛选显示下拉框（全部基因、仅显著差异基因、仅上调基因、仅下调基因）')
    add_list_item(doc, '第二列：排序方式下拉框（校正P值升序、log2倍数变化降序、log2倍数变化升序）')
    add_list_item(doc, '第三列：显示行数滑块（范围10-100，默认20）')

    add_paragraph(doc, '表格显示列包括：基因名、log2倍数变化、原始P值、校正P值、对照组均值、处理组均值、调控方向、显著性标记、预测结果。表格支持滚动查看，隐藏行索引，自动适应容器宽度。')

    add_bold_text(doc, '(4) 机器学习结果展示区')
    add_paragraph(doc, '机器学习结果区采用两列布局：')
    add_list_item(doc, '左侧：模型性能指标表格（准确率、精确率、召回率、F1分数，数值保留4位小数）')
    add_list_item(doc, '右侧：特征重要性条形图（按重要性从高到低排序）')
    add_paragraph(doc, '底部提供指标说明：准确率（预测正确的比例）、精确率（预测为显著的基因中真正显著的比例）、召回率（真正显著的基因中被成功预测的比例）、F1分数（精确率和召回率的调和平均）。')

    add_heading(doc, '3.5 下载结果页面', level=2)
    add_paragraph(doc, '下载结果页面提供所有分析结果的导出功能，分为表格数据下载区、图表下载区两个部分。')

    add_bold_text(doc, '(1) 表格数据下载区')
    add_paragraph(doc, '采用两列布局：')
    add_list_item(doc, '左侧：差异表达分析结果下载（deg_results.csv，CSV格式，UTF-8 BOM编码）')
    add_list_item(doc, '右侧：序列统计信息下载（sequence_info.csv，CSV格式）')

    add_bold_text(doc, '(2) 图表下载区')
    add_paragraph(doc, '采用两列布局，每列提供两个图表的下载：')
    add_list_item(doc, '火山图下载：volcano_plot.png，300 DPI高清')
    add_list_item(doc, '表达热图下载：expression_heatmap.png，300 DPI高清')
    add_list_item(doc, 'PCA图下载：pca_plot.png，300 DPI高清')
    add_list_item(doc, 'GC分布图下载：gc_distribution.png，300 DPI高清')

    add_paragraph(doc, '界面底部提供使用提示：所有图表均为300 DPI高分辨率PNG格式，适合用于论文和报告；CSV文件使用UTF-8 BOM编码，Excel打开不会乱码；如需其他格式或分辨率，请在【数据配置】页面调整参数后重新分析。')

    add_heading(doc, '3.6 侧边栏导航', level=2)
    add_paragraph(doc, '侧边栏作为全局导航组件，始终显示在界面左侧，提供页面切换、状态监控、版本信息展示功能。')

    add_bold_text(doc, '导航菜单区')
    add_list_item(doc, '采用单选按钮（Radio）组件形式')
    add_list_item(doc, '五个导航选项：首页、数据配置、运行分析、结果展示、下载结果')
    add_list_item(doc, '点击选项自动切换到对应页面')

    add_bold_text(doc, '状态显示区')
    add_list_item(doc, 'FASTA数据状态：已加载/未加载')
    add_list_item(doc, '表达数据状态：已加载/未加载')
    add_list_item(doc, '分析完成状态：已完成/等待分析')

    add_bold_text(doc, '版本信息区')
    add_list_item(doc, '显示平台名称与版本号："基因表达调控分析与预测平台 v1.0"')
    add_list_item(doc, '字体较小，位于侧边栏底部')

    add_heading(doc, '3.7 命令行模式操作', level=2)
    add_paragraph(doc, '除Web界面外，平台提供命令行模式运行完整分析流程，适用于批量处理、自动化脚本、服务器环境等场景。')

    add_bold_text(doc, '(1) 环境准备')
    add_paragraph(doc, '安装依赖：')
    add_code_block(doc, 'pip install -r requirements.txt')
    add_paragraph(doc, '依赖列表：streamlit>=1.28.0、numpy>=1.24.0、pandas>=2.0.0、matplotlib>=3.7.0、seaborn>=0.12.0、biopython>=1.81、scikit-learn>=1.3.0')

    add_bold_text(doc, '(2) 配置修改')
    add_paragraph(doc, '编辑config.py或创建配置文件：')
    add_code_block(doc, '''from config import AnalysisConfig

config = AnalysisConfig(
    fasta_file="data/my_genes.fasta",
    expression_file="data/my_expression.csv",
    output_dir="results",
    p_value_threshold=0.05,
    log2_fold_change_threshold=1.0,
    test_size=0.2,
    n_estimators=100
)

config.save_to_file("my_config.txt")''')

    add_bold_text(doc, '(3) 执行分析')
    add_paragraph(doc, '运行主程序：')
    add_code_block(doc, 'python main.py')
    add_paragraph(doc, '输出文件列表：')
    add_list_item(doc, 'results/deg_results.csv（差异表达结果）')
    add_list_item(doc, 'results/gc_distribution.png（GC分布图）')
    add_list_item(doc, 'results/volcano_plot.png（火山图）')
    add_list_item(doc, 'results/expression_heatmap.png（热图）')
    add_list_item(doc, 'results/pca_plot.png（PCA图）')
    add_list_item(doc, 'results/deg_results_with_classification.csv（包含预测结果）')

    add_bold_text(doc, '(4) 日志查看')
    add_paragraph(doc, '程序运行过程中在终端输出详细日志，包含时间戳、模块名、日志级别和具体信息。')

    add_bold_text(doc, '(5) 模块化调用')
    add_paragraph(doc, '用户可单独调用特定模块进行自定义分析：')
    add_code_block(doc, '''from data_processing.fasta_reader import FASTAReader
from sequence_analysis.basic_stats import SequenceStats
from expression_analysis.differential import DiffExpAnalyzer
from visualization.plots import PlotGenerator

# 读取序列
reader = FASTAReader("my_data.fasta")
sequences = reader.read_all()

# 计算GC含量
stats = SequenceStats(sequences)
gc_contents = stats.calculate_gc_content()

# 差异分析
analyzer = DiffExpAnalyzer("my_expression.csv")
deg_results = analyzer.analyze(
    group1_cols=['control_0', 'control_1'],
    group2_cols=['treatment_0', 'treatment_1']
)

# 生成图表
plotter = PlotGenerator()
plotter.plot_volcano(deg_results, output_file="my_volcano.png")''')

    doc.add_page_break()

    # ===== 第四章 系统测试 =====
    add_heading(doc, '第四章 系统测试', level=1)

    add_heading(doc, '4.1 功能测试', level=2)

    add_bold_text(doc, '(1) 数据上传模块测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '上传有效FASTA格式文件（包含10条序列）', '文件解析成功，显示序列数量10条，前3条序列ID正确显示'],
            ['TS2', '上传无效格式文件（非FASTA格式）', '提示"文件格式不正确，请上传.fasta/.fa/.txt格式文件"'],
            ['TS3', '上传表达矩阵CSV文件（100行×6列）', '文件读取成功，显示基因数量100、样本数量6、样本名称列表'],
            ['TS4', '上传缺少基因名列的CSV文件', '提示"第一列应为基因名，请检查文件格式"'],
            ['TS5', '勾选"使用示例数据"复选框', '自动加载内置数据，显示序列数量与表达矩阵维度'],
            ['TS6', '数据未上传时点击"保存配置"', '提示"请先上传FASTA序列文件"或"请先上传表达矩阵数据"'],
        ]
    )

    add_bold_text(doc, '(2) 参数配置模块测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '调整P值阈值滑块至0.01', '参数值更新为0.01，显示格式正确'],
            ['TS2', '调整log2FC阈值滑块至2.0', '参数值更新为2.0，显示格式正确'],
            ['TS3', '设置GC窗口大小为200', '输入框值更新为200，在有效范围内'],
            ['TS4', '设置测试集比例为0.3', '滑块值更新为0.3，显示正确'],
            ['TS5', '设置随机种子为2024', '输入框值更新为2024'],
            ['TS6', '设置决策树数量为200', '滑块值更新为200，范围有效'],
        ]
    )

    add_bold_text(doc, '(3) 样本分组配置测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '选择对照组3个，处理组3个，无重叠', '显示"分组配置完成：对照组3个，处理组3个"'],
            ['TS2', '两组选择相同样本', '提示"对照组和处理组不能有相同的样本"'],
            ['TS3', '对照组未选择任何样本', '提示"请至少为每组选择1个样本"'],
            ['TS4', '自动猜测对照组（列名含control）', '自动选中control_0、control_1、control_2'],
            ['TS5', '数据未上传时尝试配置分组', '提示"请先上传表达矩阵数据，再配置样本分组"'],
        ]
    )

    add_bold_text(doc, '(4) 分析流程执行测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '点击"开始完整分析流程"按钮', '进度条开始更新，日志逐条显示，四个步骤依次执行'],
            ['TS2', '观察步骤1进度（序列分析）', '进度条显示10%，日志显示GC含量计算结果'],
            ['TS3', '观察步骤2进度（差异分析）', '进度条显示35%，日志显示显著基因数量统计'],
            ['TS4', '观察步骤3进度（可视化）', '进度条显示65%，日志显示4个图表生成成功'],
            ['TS5', '观察步骤4进度（机器学习）', '进度条显示90%，日志显示模型评估指标与特征重要性'],
            ['TS6', '分析完成后查看状态', '进度条显示100%，状态文本"分析完成！"'],
            ['TS7', '点击"重新运行分析"按钮', '所有状态重置，日志清空，允许重新配置'],
        ]
    )

    add_bold_text(doc, '(5) 可视化图表显示测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '切换到"火山图"标签页', '显示火山图，红色/蓝色/灰色点区分显著基因'],
            ['TS2', '切换到"表达热图"标签页', '显示热图，红蓝颜色表示高低表达'],
            ['TS3', '切换到"PCA图"标签页', '显示PCA散点图，轴标签含方差比例'],
            ['TS4', '切换到"GC分布"标签页', '显示直方图，叠加KDE曲线，均值线与中位数线可见'],
            ['TS5', '图表未生成时查看', '显示"XX图未生成"提示信息'],
        ]
    )

    add_bold_text(doc, '(6) 结果表格显示测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '选择"仅显著差异基因"筛选', '表格仅显示显著基因，行数减少'],
            ['TS2', '选择"校正P值（升序）"排序', '表格按校正P值从小到大排列'],
            ['TS3', '调整显示行数滑块至50', '表格显示前50行数据'],
            ['TS4', '查看表格列内容', '包含基因名、log2FC、P值、校正P值、调控方向等列'],
            ['TS5', '分析未完成时查看表格', '显示"暂无差异表达分析结果"提示'],
        ]
    )

    add_bold_text(doc, '(7) 机器学习结果展示测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '查看性能指标表格', '显示准确率、精确率、召回率、F1分数'],
            ['TS2', '查看特征重要性条形图', '条形图按重要性排序'],
            ['TS3', '查看指标说明文字', '提供四项指标的公式与含义说明'],
            ['TS4', '机器学习未执行时查看', '显示"暂无机器学习分析结果"提示'],
        ]
    )

    add_bold_text(doc, '(8) 结果下载功能测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '点击"下载deg_results.csv"按钮', 'CSV文件下载成功，包含完整差异结果'],
            ['TS2', '点击"下载volcano_plot.png"按钮', 'PNG图片下载成功，300 DPI分辨率'],
            ['TS3', '点击"下载expression_heatmap.png"按钮', 'PNG图片下载成功，图表完整无缺失'],
            ['TS4', '分析未完成时尝试下载', '提示"请先在【运行分析】页面执行分析"'],
            ['TS5', '使用Excel打开CSV文件', '中文无乱码，数值格式正确'],
        ]
    )

    add_bold_text(doc, '(9) 页面导航测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '点击侧边栏"首页"选项', '页面切换到首页，显示功能概览与流程表格'],
            ['TS2', '点击侧边栏"数据配置"选项', '页面切换到数据配置，显示上传与参数设置界面'],
            ['TS3', '点击侧边栏"运行分析"选项', '页面切换到运行分析，显示配置摘要与启动按钮'],
            ['TS4', '点击侧边栏"结果展示"选项', '页面切换到结果展示，显示图表与表格'],
            ['TS5', '点击侧边栏"下载结果"选项', '页面切换到下载结果，显示下载按钮列表'],
            ['TS6', '观察侧边栏状态显示', 'FASTA/表达数据状态与分析完成状态实时更新'],
        ]
    )

    add_bold_text(doc, '(10) 命令行模式测试')
    add_table(doc,
        ['测试编号', '测试步骤', '预期结果'],
        [
            ['TS1', '执行"python main.py"命令', '程序启动，日志输出，分析流程依次执行'],
            ['TS2', '检查输出目录results/', '包含6个输出文件（CSV+PNG），文件生成成功'],
            ['TS3', '查看deg_results.csv内容', '包含差异表达结果，列结构完整，数值合理'],
            ['TS4', '打开PNG图表文件', '图表清晰，300 DPI分辨率，中文显示正常'],
            ['TS5', '修改配置参数后运行', '使用新参数重新分析，结果反映参数变化'],
            ['TS6', '模块化调用测试', '单独调用FASTAReader、SequenceStats等模块正常运行'],
        ]
    )

    add_heading(doc, '4.2 性能测试', level=2)

    add_bold_text(doc, '(1) 数据导入性能测试')
    add_paragraph(doc, '系统在完成FASTA序列文件导入（1000条序列及以内）、表达矩阵CSV文件导入（1000行×20列及以内）等数据读取操作时，响应时间控制在3秒以内，确保用户上传数据后界面即时反馈，避免长时间等待影响使用体验。大文件（>5000条序列或>5000行表达矩阵）导入时，系统采用迭代器模式逐行读取，内存占用稳定在500MB以内，不会发生内存溢出或程序崩溃。')

    add_bold_text(doc, '(2) 序列分析性能测试')
    add_paragraph(doc, '系统在执行GC含量计算、核苷酸频率统计等序列分析操作时，1000条序列处理时间控制在5秒以内，计算结果准确，日志实时输出。滑动窗口GC偏斜计算（窗口100bp）处理100条序列时间控制在10秒以内，窗口数量合理，结果分布正常。')

    add_bold_text(doc, '(3) 差异表达分析性能测试')
    add_paragraph(doc, '系统在执行t检验、FDR校正、显著性标记等差异分析操作时，1000个基因×6个样本的数据集分析时间控制在15秒以内，统计检验准确，校正算法正确执行。5000个基因规模数据集分析时间控制在60秒以内，进度条实时更新，无计算中断或超时错误。')

    add_bold_text(doc, '(4) 可视化图表生成性能测试')
    add_paragraph(doc, '系统在生成火山图、热图、PCA图、GC分布图四种图表时，单图表生成时间控制在5秒以内，图表渲染清晰，交互操作流畅。热图绘制100个基因×6个样本时间控制在10秒以内，Z-score标准化正确，聚类算法执行稳定。PCA降维计算与图表生成时间控制在8秒以内，主成分提取准确，方差比例计算正确。')

    add_bold_text(doc, '(5) 机器学习训练性能测试')
    add_paragraph(doc, '系统在执行随机森林模型训练时，1000样本×6特征数据集训练时间控制在30秒以内，模型评估指标输出完整，特征重要性计算准确。决策树数量设置为500时，训练时间控制在120秒以内，内存占用稳定，无过拟合警告。模型预测1000个新样本时间控制在5秒以内，预测结果输出及时。')

    add_bold_text(doc, '(6) 系统稳定性测试')
    add_paragraph(doc, '系统在持续运行（8小时及以上）、高频次执行分析流程（20轮及以上）、多模块协同操作（数据上传→分析→查看→下载连续操作）过程中稳定可靠，无程序崩溃、界面闪退、计算中断或无响应等异常情况。Web应用在多用户同时访问（10人并发）场景下响应时间稳定，Session State数据隔离正确，无数据混淆或状态错乱。')

    add_bold_text(doc, '(7) 数据完整性测试')
    add_paragraph(doc, '系统在执行数据上传、差异分析、图表生成、结果下载等操作时，能够保证数据的准确存储与及时更新，不会发生数据丢失、数值错乱、格式异常或导出文件内容缺失等问题。CSV导出文件UTF-8编码正确，Excel打开无乱码，数值精度保留合理。PNG图表导出300 DPI分辨率达标，中文标注清晰，图表元素完整无缺失。')

    add_bold_text(doc, '(8) 内存与资源占用测试')
    add_paragraph(doc, '系统在执行完整分析流程时，内存峰值占用控制在1GB以内，CPU利用率稳定在30%-60%范围，无异常飙升或资源耗尽现象。大文件处理时内存占用线性增长，GC回收及时，无内存泄漏。Web应用长时间运行（24小时）后内存占用保持稳定，无累积增长。')

    doc.add_page_break()

    # ===== 第五章 常见问题与解决方案 =====
    add_heading(doc, '第五章 常见问题与解决方案', level=1)

    add_heading(doc, '5.1 数据上传问题', level=2)

    add_bold_text(doc, '问题1：FASTA文件上传后显示"文件格式不正确"')
    add_paragraph(doc, '原因：文件扩展名非.fasta/.fa/.txt，或文件内部格式不符合FASTA标准（缺少">"标题行）')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '确认文件扩展名为.fasta、.fa或.txt')
    add_list_item(doc, '打开文件检查首行是否以">"开头')
    add_list_item(doc, '确认序列部分只包含ATGCN等标准碱基字符')
    add_list_item(doc, '使用文本编辑器保存为UTF-8编码格式')

    add_bold_text(doc, '问题2：表达矩阵CSV上传后部分列缺失')
    add_paragraph(doc, '原因：CSV文件分隔符错误（非逗号分隔），或列名包含特殊字符')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '使用Excel导出CSV时选择"逗号分隔"格式')
    add_list_item(doc, '列名避免使用空格、中文、特殊符号，推荐格式：gene, sample_1, sample_2')
    add_list_item(doc, '确认第一列为基因名，其他列为数值型表达量')
    add_list_item(doc, '使用记事本打开CSV检查分隔符是否为逗号')

    add_bold_text(doc, '问题3：示例数据加载失败')
    add_paragraph(doc, '原因：示例数据文件路径错误或文件损坏')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '确认项目data目录下存在simulated_genes.fasta和simulated_expression.csv文件')
    add_list_item(doc, '检查文件完整性，重新下载或生成示例数据')
    add_list_item(doc, '确认工作目录正确，避免路径混乱')

    add_heading(doc, '5.2 分析参数问题', level=2)

    add_bold_text(doc, '问题4：差异分析结果无显著基因')
    add_paragraph(doc, '原因：P值阈值过严格、log2FC阈值过大、数据本身差异小')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '在数据配置页面降低P值阈值（如从0.05调至0.1）')
    add_list_item(doc, '降低log2FC阈值（如从1.0调至0.5）')
    add_list_item(doc, '检查对照组与处理组样本分组是否正确')
    add_list_item(doc, '检查表达数据是否经过标准化处理')

    add_bold_text(doc, '问题5：机器学习模型准确率过低（<60%）')
    add_paragraph(doc, '原因：样本数太少、特征区分度不足、类别不平衡严重、超参数不合理')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '增加数据量，至少50个基因样本')
    add_list_item(doc, '检查特征列是否包含有效信息，避免全为零或常数')
    add_list_item(doc, '调整测试集比例（如从0.2调至0.1）')
    add_list_item(doc, '增加决策树数量（如从100调至200）')
    add_list_item(doc, '检查标签分布，显著基因与非显著基因比例不应过于悬殊')

    add_bold_text(doc, '问题6：GC含量计算结果异常（如0%或100%）')
    add_paragraph(doc, '原因：序列文件为空、序列仅包含N碱基、序列格式错误')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '检查FASTA文件是否包含有效序列内容')
    add_list_item(doc, '确认序列非空且长度>0')
    add_list_item(doc, '查看序列是否包含大量N（不确定碱基）')
    add_list_item(doc, '确认序列字符串已转为大写格式')

    add_heading(doc, '5.3 可视化图表问题', level=2)

    add_bold_text(doc, '问题7：图表中文显示乱码或方框')
    add_paragraph(doc, '原因：系统缺少中文字体，matplotlib默认字体不支持中文')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '确认fonts/simhei.ttf字体文件存在')
    add_list_item(doc, 'Web界面自动加载内置字体，无需手动配置')
    add_list_item(doc, '命令行模式需确保字体文件在项目根目录fonts文件夹下')
    add_list_item(doc, '如仍乱码，可手动设置matplotlib字体参数')

    add_bold_text(doc, '问题8：热图基因标签拥挤看不清')
    add_paragraph(doc, '原因：基因数量过多（>50），标签重叠')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '系统自动判断，基因>50时隐藏行标签')
    add_list_item(doc, '手动筛选显示部分基因（如仅显示显著基因）')
    add_list_item(doc, '导出高清PNG图片（300 DPI）后放大查看')
    add_list_item(doc, '在数据配置页面调整参数重新分析')

    add_bold_text(doc, '问题9：PCA图对照组与处理组未分开')
    add_paragraph(doc, '原因：样本间差异小、数据质量问题、批次效应干扰')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '检查样本分组是否正确')
    add_list_item(doc, '检查表达数据是否经过批次效应校正')
    add_list_item(doc, '确认对照组与处理组生物学差异真实存在')
    add_list_item(doc, '增加样本数量以提高统计效力')

    add_heading(doc, '5.4 结果下载问题', level=2)

    add_bold_text(doc, '问题10：CSV文件Excel打开乱码')
    add_paragraph(doc, '原因：CSV编码格式问题')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '系统默认使用UTF-8 BOM编码，Excel应正确识别')
    add_list_item(doc, '如仍乱码，使用记事本打开后另存为ANSI编码')
    add_list_item(doc, '使用LibreOffice Calc或WPS表格打开，兼容性更好')
    add_list_item(doc, '导入Excel时选择"数据→从文本/CSV"，手动指定UTF-8编码')

    add_bold_text(doc, '问题11：PNG图表下载后模糊')
    add_paragraph(doc, '原因：图片查看器缩放设置问题')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '系统导出300 DPI高清图片，分辨率达标')
    add_list_item(doc, '使用专业图片查看器（如IrfanView）打开')
    add_list_item(doc, '在图片查看器中设置100%原始尺寸显示')
    add_list_item(doc, '插入论文时保持原始分辨率，避免压缩')

    add_bold_text(doc, '问题12：下载按钮无响应')
    add_paragraph(doc, '原因：浏览器拦截下载、分析未完成')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '确认分析流程已完成，侧边栏显示"分析已完成"')
    add_list_item(doc, '检查浏览器下载设置，允许弹出窗口与下载')
    add_list_item(doc, '尝试更换浏览器（推荐Chrome、Edge、Firefox）')
    add_list_item(doc, '清除浏览器缓存后重新尝试')

    add_heading(doc, '5.5 命令行模式问题', level=2)

    add_bold_text(doc, '问题13：依赖安装失败')
    add_paragraph(doc, '原因：Python版本过低、pip版本过旧、网络问题')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '确认Python版本>=3.8，推荐3.10+')
    add_list_item(doc, '升级pip：python -m pip install --upgrade pip')
    add_list_item(doc, '使用国内镜像源加速下载')
    add_list_item(doc, '分步安装依赖，逐个排查问题包')

    add_bold_text(doc, '问题14：main.py执行报错"找不到模块"')
    add_paragraph(doc, '原因：工作目录错误、模块未导入、路径问题')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '确认在项目根目录执行命令：cd GeneAnalysis')
    add_list_item(doc, '检查sys.path是否包含项目路径')
    add_list_item(doc, '确认各模块__init__.py文件存在')
    add_list_item(doc, '可使用绝对路径导入方式')

    add_bold_text(doc, '问题15：输出目录results/未生成文件')
    add_paragraph(doc, '原因：分析流程中途错误、权限问题、路径问题')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '查看终端日志定位错误位置')
    add_list_item(doc, '确认有写入权限')
    add_list_item(doc, '手动创建输出目录：mkdir results')
    add_list_item(doc, '检查config.py中output_dir配置正确')

    add_heading(doc, '5.6 Web应用问题', level=2)

    add_bold_text(doc, '问题16：Streamlit应用首次加载慢')
    add_paragraph(doc, '原因：首次运行需安装依赖、下载字体')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '等待1-3分钟，依赖安装完成后后续访问快速')
    add_list_item(doc, '云端部署版本首次加载较慢，属于正常现象')
    add_list_item(doc, '本地运行预先安装依赖：pip install -r requirements.txt')
    add_list_item(doc, '避免频繁刷新页面，减少重新加载次数')

    add_bold_text(doc, '问题17：Session State数据丢失')
    add_paragraph(doc, '原因：页面刷新、浏览器关闭、Session超时')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '避免刷新页面，使用侧边栏导航切换')
    add_list_item(doc, '分析完成后及时下载结果保存本地')
    add_list_item(doc, '重要配置保存到文件：config.save_to_file("my_config.txt")')
    add_list_item(doc, '命令行模式无Session限制，适合长时间分析')

    add_bold_text(doc, '问题18：侧边栏状态显示不准确')
    add_paragraph(doc, '原因：Session State更新延迟、分析中断')
    add_paragraph(doc, '解决方案：')
    add_list_item(doc, '等待分析流程完全完成再查看状态')
    add_list_item(doc, '如显示异常，点击"重新运行分析"重置状态')
    add_list_item(doc, '查看运行分析页面日志确认实际进度')
    add_list_item(doc, '刷新页面（F5）强制更新Session State')

    doc.add_page_break()

    # ===== 第六章 附录 =====
    add_heading(doc, '第六章 附录', level=1)

    add_heading(doc, '6.1 示例数据格式说明', level=2)

    add_bold_text(doc, 'FASTA序列文件格式示例：')
    add_code_block(doc, '''>gene_001 Homo sapiens TP53 tumor protein p53
ATGAGCCACCCTGAGCCGGCTCCTGATTCCTTTCTTT...
>gene_002 Homo sapiens BRCA1 breast cancer 1
ATGCGATCGATCGATCGATCGATCGATCG...
>gene_003 Homo sapiens GATA1 GATA-binding protein 1
ATGCGATCGATCGATCGATCG...''')

    add_paragraph(doc, '格式要求：')
    add_list_item(doc, '以">"开头的行为标题行，包含基因ID与描述信息')
    add_list_item(doc, '标题行后为序列内容，仅包含ATGCN标准碱基字符')
    add_list_item(doc, '序列可跨多行书写，系统自动合并')
    add_list_item(doc, '基因ID建议使用简洁命名（如gene_001），避免空格与特殊字符')

    add_bold_text(doc, '表达矩阵CSV文件格式示例：')
    add_code_block(doc, '''gene,control_0,control_1,control_2,treatment_0,treatment_1,treatment_2
gene_001,100,120,110,50,60,55
gene_002,50,55,48,200,210,195
gene_003,80,85,78,120,125,118''')

    add_paragraph(doc, '格式要求：')
    add_list_item(doc, '第一列为基因名，与FASTA文件基因ID对应')
    add_list_item(doc, '其他列为样本表达量，数值型数据')
    add_list_item(doc, '列名为样本名，建议包含分组标识（如control/treatment）')
    add_list_item(doc, '每行一个基因，无空行与注释行')
    add_list_item(doc, '文件保存为逗号分隔CSV格式，UTF-8编码')

    add_heading(doc, '6.2 参数调整建议参考表', level=2)
    add_table(doc,
        ['参数名称', '默认值', '探索性研究建议', '验证性研究建议', '数据量少建议'],
        [
            ['p_value_threshold', '0.05', '0.10（放宽）', '0.01（严格）', '0.05（标准）'],
            ['log2_fold_change_threshold', '1.0', '0.5（降低）', '1.5-2.0（提高）', '0.5-1.0'],
            ['test_size', '0.2', '0.3（增加测试）', '0.15（减少测试）', '0.1-0.15'],
            ['n_estimators', '100', '100-200', '200-500', '50-100'],
            ['gc_window_size', '100', '50（高分辨）', '200-500（低分辨）', '100'],
            ['random_state', '42', '固定不变', '固定不变', '固定不变'],
        ]
    )

    add_heading(doc, '6.3 技术支持与联系方式', level=2)

    add_bold_text(doc, '在线文档：')
    add_list_item(doc, '项目README：README.md')
    add_list_item(doc, '代码Wiki：CODE_WIKI.md')
    add_list_item(doc, '模块规划：module_planning.md')

    add_bold_text(doc, '在线应用地址：')
    add_list_item(doc, 'Streamlit Cloud：https://geneanalysis-app.streamlit.app/')

    add_bold_text(doc, '问题反馈渠道：')
    add_list_item(doc, 'GitHub Issues：提交Bug报告与功能建议')
    add_list_item(doc, 'Pull Requests：贡献代码改进')

    add_bold_text(doc, '开发团队：')
    add_list_item(doc, 'GeneAnalysis Team')

    doc.add_paragraph()
    doc.add_paragraph()

    # 文档版本信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 文档结束 —')
    set_chinese_font(run, "宋体", 12)
    run.italic = True

    doc.add_paragraph()
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run('文档名称：基因表达调控分析与预测平台（V1.0）软件使用说明书\n文档版本：V1.0\n编制日期：2024年1月\n编制单位：GeneAnalysis Team')
    set_chinese_font(run, "宋体", 10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 保存文档
    output_path = '/workspace/基因表达调控分析与预测平台（V1.0）软件使用说明书.docx'
    doc.save(output_path)
    print(f'Word文档已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    generate_manual()
